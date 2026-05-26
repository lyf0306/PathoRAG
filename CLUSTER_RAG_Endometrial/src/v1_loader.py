import os
from pathlib import Path
from typing import List, Dict
import json
from datetime import datetime

import pdfplumber
from docx import Document


def extract_text_from_pdf(pdf_path: Path) -> str:
    """从 PDF 中提取文本（包括普通文本和表格内容）"""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            all_text_parts = []
            # 普通页面文本
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    all_text_parts.append(page_text)
            # 表格内容（按行输出，单元格用 | 分隔）
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if any(cell and str(cell).strip() for cell in row):
                            row_text = " | ".join(
                                str(cell).strip() if cell else "" for cell in row
                            )
                            all_text_parts.append(row_text)
            return "\n".join(all_text_parts).strip()
    except Exception as e:
        print(f"⚠️ PDF 读取失败 [{pdf_path.name}]: {e}")
        return ""


def extract_text_from_docx(docx_path: Path) -> str:
    """
    从 Word 文档中提取文本（包括段落和所有表格，支持嵌套表格）。
    自动去除因合并单元格导致的重复文本。
    """
    try:
        doc = Document(docx_path)
        all_text_parts = []

        # 辅助函数：递归提取单元格内的所有文本（段落 + 嵌套表格）
        def extract_cell_text(cell):
            texts = []
            # 单元格内的段落
            for para in cell.paragraphs:
                if para.text.strip():
                    texts.append(para.text.strip())
            # 单元格内的嵌套表格
            for sub_table in cell.tables:
                for row in sub_table.rows:
                    for sub_cell in row.cells:
                        texts.extend(extract_cell_text(sub_cell))
            return texts

        # 1. 普通段落
        for para in doc.paragraphs:
            if para.text.strip():
                all_text_parts.append(para.text.strip())

        # 2. 顶层表格（处理合并单元格导致的重复）
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                prev_cell_text = None  # 记录上一个单元格的文本，用于去重
                for cell in row.cells:
                    cell_content = extract_cell_text(cell)
                    if cell_content:
                        cell_text = " ".join(cell_content)
                        # 去重：如果当前单元格文本与上一个单元格文本相同，则跳过（合并单元格导致）
                        if cell_text != prev_cell_text:
                            row_texts.append(cell_text)
                            prev_cell_text = cell_text
                    else:
                        prev_cell_text = None  # 空单元格重置，避免影响
                if row_texts:
                    all_text_parts.append(" | ".join(row_texts))

        return "\n".join(all_text_parts).strip()
    except Exception as e:
        print(f"⚠️ DOCX 读取失败 [{docx_path.name}]: {e}")
        return ""


def process_folder(folder_path: str, recursive: bool = False) -> List[Dict[str, str]]:
    """遍历文件夹，提取所有 PDF 和 DOCX 文件的文本内容。"""
    results = []
    base_dir = Path(folder_path)

    if not base_dir.is_dir():
        raise NotADirectoryError(f"路径不存在或不是文件夹: {folder_path}")

    iterator = base_dir.rglob("*") if recursive else base_dir.iterdir()

    for file_path in iterator:
        if not file_path.is_file():
            continue

        ext = file_path.suffix.lower()
        if ext == ".pdf":
            text = extract_text_from_pdf(file_path)
            file_type = "PDF"
        elif ext == ".docx":
            text = extract_text_from_docx(file_path)
            file_type = "DOCX"
        else:
            continue

        status = "成功" if text else "警告：提取内容为空"
        results.append({
            "id": file_path.name,
            "file_path": str(file_path),
            "type": file_type,
            "raw_text": text,
            "text_length": len(text),
            "status": status
        })

    return results


def print_detailed_report(data: List[Dict[str, str]], folder_path: str, recursive: bool):
    """打印详细的处理报告到控制台，并保存为 JSON 和日志文件"""
    print("\n" + "=" * 80)
    print(f"📁 源文件夹: {folder_path}  (递归模式: {'是' if recursive else '否'})")
    print(f"📄 共处理文件数: {len(data)}")
    print("=" * 80)

    success_count = sum(1 for item in data if item["status"] == "成功")
    warn_count = len(data) - success_count
    total_length = sum(item["text_length"] for item in data)

    print(f"✅ 成功提取: {success_count} 个")
    print(f"⚠️  空内容警告: {warn_count} 个")
    print(f"📊 总文本字符数: {total_length:,}")
    print("-" * 80)

    for idx, item in enumerate(data, 1):
        print(f"\n[{idx}] 文件: {item['id']}")
        print(f"    类型: {item['type']}")
        print(f"    状态: {item['status']}")
        print(f"    文本长度: {item['text_length']} 字符")
        preview = item['raw_text'][:200].replace('\n', ' ')
        if len(item['raw_text']) > 200:
            preview += "..."
        print(f"    预览: {preview}")
        if item['status'] != "成功":
            print(f"    ⚠️  请检查文件是否包含可提取的文本（可能是扫描件或保护文档）")

    print("\n" + "=" * 80)
    print("✅ 处理完成")
    print("=" * 80)

    # 保存完整结果到 JSON
    output_json = Path(folder_path).parent / "extracted_output.json"
    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"\n💾 完整结果已保存至: {output_json}")

    # 保存日志文件
    log_file = Path(folder_path).parent / f"extraction_log_{datetime.now():%Y%m%d_%H%M%S}.txt"
    with open(log_file, "w", encoding="utf-8") as f:
        f.write(f"Extraction Report - {datetime.now()}\n")
        f.write(f"Folder: {folder_path}\n")
        f.write(f"Recursive: {recursive}\n")
        f.write(f"Total files: {len(data)}\n")
        f.write(f"Successful: {success_count}, Warnings: {warn_count}\n")
        f.write("-" * 80 + "\n")
        for item in data:
            f.write(f"File: {item['id']}\n")
            f.write(f"Status: {item['status']}\n")
            f.write(f"Length: {item['text_length']}\n")
            f.write(f"Preview: {item['raw_text'][:200]}\n")
            f.write("\n")
    print(f"📝 日志文件已保存至: {log_file}")


if __name__ == "__main__":
    # ========== 请修改为你的实际路径 ==========
    target_folder = "../data/ori"   # 你的数据文件夹路径
    recursive_mode = False          # 是否递归子文件夹
    # ========================================

    try:
        extracted_data = process_folder(target_folder, recursive=recursive_mode)
        print_detailed_report(extracted_data, target_folder, recursive_mode)
    except Exception as e:
        print(f"❌ 程序执行出错: {e}")