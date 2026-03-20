import os
import re
import json
import asyncio
from docx import Document
import fitz  # PyMuPDF，用于处理 PDF
from openai import AsyncOpenAI
from tqdm.asyncio import tqdm

# ================= 配置区域 =================
INPUT_DIR = "/root/data_for_patients"
OUTPUT_DIR = "/root/result"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "clinical_dataset_v3.jsonl") # 升级为 v3

DEEPSEEK_API_URL = "https://api.deepseek.com/v1" 
DEEPSEEK_API_KEY = "sk-90d43b1588094b2ba97a4e75f845134a" 
LLM_MODEL_NAME = "deepseek-chat" 
# ===========================================

os.makedirs(OUTPUT_DIR, exist_ok=True)

def extract_text_from_file(filepath):
    """从 Word 或 PDF 文件中提取纯文本"""
    text = ""
    try:
        if filepath.endswith('.docx'):
            doc = Document(filepath)
            text = "\n".join([para.text for para in doc.paragraphs])
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + " | ".join([cell.text for cell in row.cells])
        elif filepath.endswith('.pdf'):
            with fitz.open(filepath) as doc:
                for page in doc:
                    text += page.get_text() + "\n"
    except Exception as e:
        print(f"读取文件 {filepath} 失败: {e}")
    return text

async def distill_with_llm(raw_text, filename, client):
    """调用 DeepSeek，执行全景特征提取与复杂性判定"""
    
    # 简单清洗多余空格和换行，控制输入 Token
    clean_text = re.sub(r'\s+', ' ', raw_text).strip()[:4000]
    if len(clean_text) < 50:
        return None
        
    system_prompt = (
        "你是一位资深的妇科肿瘤临床数据专家。你的任务是从带有大量口语化和冗余信息的原始病历中，提取核心临床特征，并将其转化为高度浓缩、结构化的医学实体标签序列。\n\n"
        "【提取规则】\n"
        "1. **绝对剔除噪音**：忽略所有就诊过程（如“当地医院就诊”）、患者行为（如“未曾体检”）、口语化症状描述，只提取标准的医学术语。\n"
        "2. **术前特征提取**：\n"
        "   - 基础状态：年龄、绝经状态（如“绝经10年”）。\n"
        "   - 核心症状：如“绝经后阴道流血”。\n"
        "   - 关键检查：肿瘤标志物异常（如“CA125升高”）、超声/核磁关键阳性发现（如“累及深肌层”）。\n"
        "   - 术前病理：如宫腔镜活检提示的病理类型。\n"
        "3. **术后病理提取**：\n"
        "   - 最终病理诊断、组织学分级(G1/G2/G3)、肌层浸润深度、脉管癌栓(LVSI)、淋巴结状态、免疫组化(如pMMR/dMMR)、FIGO分期等。\n"
        "4. **复杂病例判定 (MoE 路由标签)**：\n"
        "   - 如果患者患有高危合并症（如肥胖、高血压、糖尿病、脂肪肝等），或存在特殊基因突变/晚期转移（III/IV期），请必须将其判定为复杂病例。\n\n"
        "【输出格式】\n"
        "将提取出的特征用逗号分隔，按“患者状态 -> 术前检查 -> 术后病理 -> 合并症”的逻辑顺序排列。\n"
        "同时，从文本中摘录出具体的放化疗及随访方案，作为【真实方案】。\n"
        "必须严格按照以下格式输出，使用 <keywords> 标签包裹最终结果：\n"
        "<think>\n逐步分析过程...\n</think>\n"
        "<keywords>\n"
        "[中文检索词]: 59岁, 绝经10年, 绝经后阴道流血, CA125升高, 宫腔肿物累及深肌层, 术前非典型增生, 子宫内膜样癌, Ⅰ级, 复杂不典型增生, FIGO IA期, 2型糖尿病\n"
        "[英文检索词]: 59 years old, 10 years postmenopause, Postmenopausal vaginal bleeding, Elevated CA125, Intrauterine mass involving deep myometrium, Preoperative atypical hyperplasia, Endometrioid adenocarcinoma, Grade 1, Complex atypical hyperplasia, FIGO stage IA, Type 2 diabetes\n"
        "[真实方案]: 腹腔镜全子宫切除+双附件切除；术后辅助TC方案化疗；前2年每3个月随访1次。\n"
        "[是否复杂病例]: True\n"
        "</keywords>"
    )
    
    user_prompt = f"【原始全景病历数据】\n{clean_text}\n"

    try:
        response = await client.chat.completions.create(
            model=LLM_MODEL_NAME,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.1,
            max_tokens=2048
        )
        
        result_str = response.choices[0].message.content.strip()
        
        keywords_match = re.search(r'<keywords>([\s\S]*?)</keywords>', result_str)
        if not keywords_match:
            print(f"\n[Warning] {filename} 未检测到标准输出格式！")
            return None
            
        core_content = keywords_match.group(1).strip()
        
        cn_keywords = re.search(r'\[中文检索词\]:\s*(.*)', core_content)
        en_keywords = re.search(r'\[英文检索词\]:\s*(.*)', core_content)
        raw_treatment = re.search(r'\[真实方案\]:\s*(.*)', core_content)
        is_complex_match = re.search(r'\[是否复杂病例\]:\s*(True|False|true|false)', core_content)
        
        is_complex_flag = False
        if is_complex_match and is_complex_match.group(1).lower() == 'true':
            is_complex_flag = True
            
        result_json = {
            "source_file": filename,
            "anchor_cn": cn_keywords.group(1).strip() if cn_keywords else "",
            "anchor_en": en_keywords.group(1).strip() if en_keywords else "",
            "raw_positive": raw_treatment.group(1).strip() if raw_treatment else "",
            "is_complex": is_complex_flag
        }
        
        return result_json
        
    except Exception as e:
        print(f"\n[Warning] {filename} DeepSeek API 调用或解析异常: {e}")
        return None

async def process_file(filepath, filename, client):
    raw_text = extract_text_from_file(filepath)
    distilled_data = await distill_with_llm(raw_text, filename, client)
    return distilled_data

async def main():
    print(f">>> 开始扫描 {INPUT_DIR} 目录下的患者文档...")
    valid_files = [f for f in os.listdir(INPUT_DIR) if (f.endswith(".docx") or f.endswith(".pdf")) and not f.startswith("~$")]
    print(f"找到 {len(valid_files)} 个有效文档，准备启动提炼流水线...\n")
    
    llm_client = AsyncOpenAI(base_url=DEEPSEEK_API_URL, api_key=DEEPSEEK_API_KEY)
    results = []
    semaphore = asyncio.Semaphore(10)
    
    async def bound_process(filename):
        filepath = os.path.join(INPUT_DIR, filename)
        async with semaphore:
            return await process_file(filepath, filename, llm_client)

    tasks = [bound_process(f) for f in valid_files]
    for completed_task in tqdm(asyncio.as_completed(tasks), total=len(tasks), desc="数据提炼进度"):
        result = await completed_task
        if result and result["anchor_cn"]: # 确保提炼出了内容
            results.append(result)
            
    print(f"\n>>> 提炼完成！成功获取 {len(results)} 条全景结构化数据。")
    with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
        for item in results:
            f.write(json.dumps(item, ensure_ascii=False) + '\n')
    print("✅ 保存成功！")

if __name__ == "__main__":
    asyncio.run(main())